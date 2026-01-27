"""
Извлечение метаданных из DOCX документов СМК
Более точный источник для названий и структурированных данных
"""

import re
from pathlib import Path
from typing import Optional, Dict, List
from dataclasses import dataclass, field


@dataclass
class DocxMetadata:
    """Метаданные из DOCX файла"""
    title: Optional[str] = None           # Полное название
    doc_number: Optional[str] = None      # Номер документа
    doc_type: Optional[str] = None        # Тип документа
    department: Optional[str] = None      # Подразделение
    developer: Optional[str] = None       # Разработчик
    approval_date: Optional[str] = None   # Дата утверждения
    effective_date: Optional[str] = None  # Дата введения


# Ключевые слова для поиска названия в таблицах
TITLE_KEYWORDS = [
    'наименование документа',
    'название документа',
    'document name',
    'наименование',
    'название',
]

# Ключевые слова для других полей
FIELD_KEYWORDS = {
    'doc_number': ['номер документа', 'document number', 'регистрационный номер'],
    'doc_type': ['вид документа', 'тип документа', 'document type'],
    'department': ['подразделение', 'department', 'отдел'],
    'developer': ['разработчик', 'developer', 'автор'],
    'approval_date': ['дата утверждения', 'approval date', 'утвержден'],
    'effective_date': ['дата введения', 'effective date', 'дата ввода'],
}


def extract_from_docx(docx_path: Path) -> DocxMetadata:
    """
    Извлечь метаданные из DOCX файла
    
    Стратегия поиска:
    1. Ищем таблицу со структурой "Ключ | Значение" (обычно Таблица 1)
       Строка "Наименование документа" содержит название
    2. Если не найдено - ищем в таблице на титульной странице
    3. Проверяем свойства документа
    """
    try:
        from docx import Document
    except ImportError:
        return DocxMetadata()
    
    metadata = DocxMetadata()
    
    try:
        doc = Document(docx_path)
    except Exception as e:
        return DocxMetadata()
    
    # 1. Ищем таблицу с метаданными (структура "Ключ | Значение")
    for table in doc.tables[:10]:  # Проверяем первые 10 таблиц
        # Проверяем, похожа ли таблица на таблицу метаданных (2-3 столбца)
        if len(table.columns) < 2 or len(table.columns) > 4:
            continue
        
        for row in table.rows:
            cells = row.cells
            if len(cells) < 2:
                continue
            
            # Берём первую и вторую ячейку как ключ-значение
            key_text = clean_text(cells[0].text).lower()
            value_text = clean_text(cells[1].text)
            
            # Пропускаем пустые строки
            if not key_text or not value_text:
                continue
            
            # Пропускаем если значение слишком короткое
            if len(value_text) < 5:
                continue
            
            # Название документа
            if any(kw in key_text for kw in TITLE_KEYWORDS):
                if len(value_text) > 10 and not metadata.title:
                    metadata.title = value_text
            
            # Номер документа
            if any(kw in key_text for kw in FIELD_KEYWORDS['doc_number']):
                if not metadata.doc_number:
                    metadata.doc_number = value_text
            
            # Тип документа
            if any(kw in key_text for kw in FIELD_KEYWORDS['doc_type']):
                if not metadata.doc_type:
                    metadata.doc_type = value_text
            
            # Подразделение
            if any(kw in key_text for kw in FIELD_KEYWORDS['department']):
                if not metadata.department:
                    metadata.department = value_text
            
            # Разработчик
            if any(kw in key_text for kw in FIELD_KEYWORDS['developer']):
                if not metadata.developer:
                    metadata.developer = value_text
    
    # 2. Если название не найдено, пробуем таблицу титульной страницы
    if not metadata.title and doc.tables:
        title = extract_title_from_title_page(doc.tables[0])
        if title:
            metadata.title = title
    
    # 3. Проверяем свойства документа
    if not metadata.title and doc.core_properties.title:
        metadata.title = clean_text(doc.core_properties.title)
    
    return metadata


def extract_title_from_title_page(table) -> Optional[str]:
    """
    Извлечь название с титульной страницы (первая таблица)
    
    Обычно название находится в одной из первых строк,
    в ячейке с большим количеством текста
    """
    candidates = []
    
    for row in table.rows[:5]:  # Первые 5 строк
        for cell in row.cells:
            text = clean_text(cell.text)
            
            # Пропускаем служебные поля
            skip_patterns = [
                'дата введения', 'effective date',
                'утвержден', 'approved',
                'система менеджмента', 'quality management',
                'версия', 'version', 'revision',
                'страница', 'page',
            ]
            
            text_lower = text.lower()
            if any(pat in text_lower for pat in skip_patterns):
                continue
            
            # Кандидат должен быть достаточно длинным
            if len(text) > 20 and len(text) < 200:
                # Не должен содержать много цифр (это не номер)
                digit_ratio = sum(c.isdigit() for c in text) / len(text)
                if digit_ratio < 0.3:
                    candidates.append(text)
    
    # Выбираем самый длинный осмысленный текст
    if candidates:
        return max(candidates, key=len)
    
    return None


def clean_text(text: str) -> str:
    """Очистка текста от лишних символов"""
    if not text:
        return ""
    
    # Убираем переносы строк
    text = text.replace('\n', ' ').replace('\r', ' ')
    
    # Убираем множественные пробелы
    text = re.sub(r'\s+', ' ', text)
    
    # Убираем подчеркивания (часто используются как линии)
    text = re.sub(r'_+', '', text)
    
    # Убираем начальные/конечные пробелы
    text = text.strip()
    
    return text


def find_docx_for_pdf(pdf_path: Path, docx_base_path: Path = None) -> Optional[Path]:
    """
    Найти соответствующий DOCX файл для PDF
    
    Структура:
    - pdf/КОД ^UNID/КОД (Эталон для печати).pdf
    - docx/КОД ^UNID/КОД версия N.docx
    """
    if docx_base_path is None:
        # Предполагаем структуру input2/BND/
        docx_base_path = pdf_path.parent.parent.parent / "docx"
    
    if not docx_base_path.exists():
        return None
    
    # Имя папки с PDF (содержит код и UNID)
    folder_name = pdf_path.parent.name
    
    # Ищем такую же папку в docx
    docx_folder = docx_base_path / folder_name
    
    if docx_folder.exists():
        # Ищем docx файлы
        docx_files = list(docx_folder.glob("*.docx")) + list(docx_folder.glob("*.doc"))
        if docx_files:
            return docx_files[0]
    
    return None


def batch_extract_docx_metadata(pdf_paths: List[Path], docx_base_path: Path = None) -> Dict[str, DocxMetadata]:
    """
    Пакетное извлечение метаданных из DOCX для списка PDF
    
    Returns:
        Dict[doc_code -> DocxMetadata]
    """
    from .parser import parse_document_code
    
    results = {}
    
    for pdf_path in pdf_paths:
        # Получаем код документа
        doc = parse_document_code(pdf_path.parent.name if pdf_path.parent.name != 'pdf' else pdf_path.name)
        if not doc:
            continue
        
        # Ищем соответствующий docx
        docx_path = find_docx_for_pdf(pdf_path, docx_base_path)
        if docx_path:
            metadata = extract_from_docx(docx_path)
            results[doc.code] = metadata
    
    return results


if __name__ == "__main__":
    # Тест
    import sys
    
    test_path = Path("/home/budnik_an/Obligations/input2/BND/docx/ДП-Б1.004-06 ^8CD15E6999417E4045258AF40040636E/ДП-Б1.004-06 версия 1.docx")
    
    if len(sys.argv) > 1:
        test_path = Path(sys.argv[1])
    
    if test_path.exists():
        print(f"📄 Анализ: {test_path.name}")
        print("=" * 60)
        
        metadata = extract_from_docx(test_path)
        
        print(f"📖 Название: {metadata.title or 'Не найдено'}")
        print(f"🔢 Номер: {metadata.doc_number or 'Не найден'}")
        print(f"📂 Тип: {metadata.doc_type or 'Не найден'}")
        print(f"🏢 Подразделение: {metadata.department or 'Не найдено'}")
        print(f"👤 Разработчик: {metadata.developer or 'Не найден'}")
        print(f"📅 Дата введения: {metadata.effective_date or 'Не найдена'}")
    else:
        print(f"❌ Файл не найден: {test_path}")
