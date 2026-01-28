"""
DOCX Parser - извлечение структуры из Word документов

Функции:
- Извлечение заголовков (Heading 1, 2, 3...)
- Валидация соответствия DOCX и PDF
- Извлечение метаданных
"""

import re
from pathlib import Path
from typing import List, Optional, Tuple, Dict
from dataclasses import dataclass

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    FITZ_AVAILABLE = True
except ImportError:
    FITZ_AVAILABLE = False


@dataclass
class Heading:
    """Заголовок документа"""
    text: str
    level: int  # 1, 2, 3...
    
    def __str__(self):
        return f"{'#' * self.level} {self.text}"


@dataclass
class ValidationResult:
    """Результат валидации DOCX vs PDF"""
    is_valid: bool
    pages_match: bool
    doc_code_match: bool
    last_page_match: bool
    details: str
    
    pdf_pages: int = 0
    docx_pages: int = 0
    pdf_doc_code: str = ""
    docx_doc_code: str = ""


# Regex для извлечения кода документа
DOC_CODE_PATTERN = re.compile(
    r'(КД|ДП|РД|РИ|СТ|РГ|ИОТ|TPM|ПР)-[А-ЯA-Z0-9\.\-]+',
    re.IGNORECASE
)


def extract_doc_code(text: str) -> Optional[str]:
    """Извлечь код документа из текста"""
    if not text:
        return None
    match = DOC_CODE_PATTERN.search(text)
    return match.group(0) if match else None


def extract_structure_docx(docx_path: str) -> List[Heading]:
    """
    Извлечь структуру заголовков из DOCX
    
    Args:
        docx_path: Путь к DOCX файлу
        
    Returns:
        Список заголовков с уровнями
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx не установлен. Установите: pip install python-docx")
    
    doc = Document(docx_path)
    headings = []
    
    for p in doc.paragraphs:
        if p.style and 'Heading' in p.style.name:
            text = p.text.strip()
            if text:  # Пропускаем пустые заголовки
                # Извлекаем уровень из названия стиля
                try:
                    level = int(p.style.name.replace('Heading ', '').replace('Heading', '1'))
                except ValueError:
                    level = 1
                headings.append(Heading(text=text, level=level))
    
    return headings


def get_docx_last_content(docx_path: str) -> str:
    """Получить текст последних параграфов DOCX"""
    if not DOCX_AVAILABLE:
        return ""
    
    doc = Document(docx_path)
    
    # Собираем текст последних непустых параграфов
    last_texts = []
    for p in reversed(doc.paragraphs):
        text = p.text.strip()
        if text:
            last_texts.insert(0, text)
            if len(last_texts) >= 3:
                break
    
    return ' '.join(last_texts)


def get_docx_first_content(docx_path: str) -> str:
    """Получить текст первых параграфов DOCX"""
    if not DOCX_AVAILABLE:
        return ""
    
    doc = Document(docx_path)
    
    # Собираем текст первых непустых параграфов
    first_texts = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            first_texts.append(text)
            if len(first_texts) >= 5:
                break
    
    return ' '.join(first_texts)


def estimate_docx_pages(docx_path: str) -> int:
    """
    Оценить количество страниц в DOCX
    
    Точный подсчёт без Word невозможен, используем эвристику
    """
    if not DOCX_AVAILABLE:
        return 0
    
    doc = Document(docx_path)
    
    # Способ 1: По количеству секций (если есть page breaks)
    section_count = len(doc.sections)
    
    # Способ 2: Считаем page breaks в документе
    page_breaks = 0
    for p in doc.paragraphs:
        for run in p.runs:
            if run._element.xml and 'w:br' in run._element.xml and 'w:type="page"' in run._element.xml:
                page_breaks += 1
    
    # Способ 3: Эвристика по объёму текста (~3000 символов на страницу)
    total_chars = sum(len(p.text) for p in doc.paragraphs)
    estimated_by_chars = max(1, total_chars // 3000)
    
    # Берём максимум из оценок
    return max(section_count, page_breaks + 1, estimated_by_chars)


def get_pdf_pages(pdf_path: str) -> int:
    """Получить количество страниц в PDF"""
    if not FITZ_AVAILABLE:
        return 0
    
    with fitz.open(pdf_path) as doc:
        return doc.page_count


def get_pdf_first_page_text(pdf_path: str) -> str:
    """Получить текст первой страницы PDF"""
    if not FITZ_AVAILABLE:
        return ""
    
    with fitz.open(pdf_path) as doc:
        if doc.page_count > 0:
            return doc[0].get_text()
    return ""


def get_pdf_last_page_text(pdf_path: str) -> str:
    """Получить текст последней страницы PDF"""
    if not FITZ_AVAILABLE:
        return ""
    
    with fitz.open(pdf_path) as doc:
        if doc.page_count > 0:
            return doc[-1].get_text()
    return ""


def text_similarity(text1: str, text2: str) -> float:
    """
    Простая мера похожести текстов
    
    Используем Jaccard similarity по словам
    """
    if not text1 or not text2:
        return 0.0
    
    words1 = set(text1.lower().split())
    words2 = set(text2.lower().split())
    
    if not words1 or not words2:
        return 0.0
    
    intersection = words1 & words2
    union = words1 | words2
    
    return len(intersection) / len(union) if union else 0.0


def validate_docx_vs_pdf(docx_path: str, pdf_path: str) -> ValidationResult:
    """
    Проверить соответствие DOCX и PDF
    
    Критерии:
    1. Количество страниц (±1)
    2. Код документа на первой странице
    3. Контент последней страницы (начало и конец)
    
    Args:
        docx_path: Путь к DOCX
        pdf_path: Путь к PDF
        
    Returns:
        ValidationResult с деталями проверки
    """
    details = []
    
    # 1. Количество страниц
    pdf_pages = get_pdf_pages(pdf_path)
    docx_pages = estimate_docx_pages(docx_path)
    pages_match = abs(pdf_pages - docx_pages) <= 1
    
    details.append(f"Страницы: PDF={pdf_pages}, DOCX≈{docx_pages}, {'OK' if pages_match else 'MISMATCH'}")
    
    # 2. Код документа
    pdf_first = get_pdf_first_page_text(pdf_path)
    docx_first = get_docx_first_content(docx_path)
    
    pdf_doc_code = extract_doc_code(pdf_first) or ""
    docx_doc_code = extract_doc_code(docx_first) or ""
    
    doc_code_match = (
        pdf_doc_code.upper() == docx_doc_code.upper() 
        if pdf_doc_code and docx_doc_code 
        else True  # Если код не найден - не считаем ошибкой
    )
    
    details.append(f"Код: PDF='{pdf_doc_code}', DOCX='{docx_doc_code}', {'OK' if doc_code_match else 'MISMATCH'}")
    
    # 3. Контент последней страницы
    pdf_last = get_pdf_last_page_text(pdf_path)
    docx_last = get_docx_last_content(docx_path)
    
    # Берём начало и конец для сравнения
    pdf_snippet = (pdf_last[:100] + pdf_last[-100:]) if len(pdf_last) > 200 else pdf_last
    docx_snippet = (docx_last[:100] + docx_last[-100:]) if len(docx_last) > 200 else docx_last
    
    similarity = text_similarity(pdf_snippet, docx_snippet)
    last_page_match = similarity >= 0.5  # Порог 50%
    
    details.append(f"Последняя стр.: similarity={similarity:.2f}, {'OK' if last_page_match else 'MISMATCH'}")
    
    # Итоговый вердикт
    is_valid = pages_match and doc_code_match and last_page_match
    
    return ValidationResult(
        is_valid=is_valid,
        pages_match=pages_match,
        doc_code_match=doc_code_match,
        last_page_match=last_page_match,
        details='; '.join(details),
        pdf_pages=pdf_pages,
        docx_pages=docx_pages,
        pdf_doc_code=pdf_doc_code,
        docx_doc_code=docx_doc_code,
    )


def find_docx_for_pdf(pdf_path: str, docx_base_dir: str = None) -> Optional[str]:
    """
    Найти DOCX файл, соответствующий PDF
    
    Логика поиска:
    1. В той же папке с тем же именем
    2. В параллельной папке /docx/ (если PDF в /pdf/)
    3. По коду документа в имени файла
    
    Args:
        pdf_path: Путь к PDF
        docx_base_dir: Базовая директория для поиска DOCX
        
    Returns:
        Путь к DOCX или None
    """
    pdf_path = Path(pdf_path)
    
    # Извлекаем код документа из имени файла
    doc_code = extract_doc_code(pdf_path.stem)
    
    # Способ 1: Параллельная структура /pdf/ -> /docx/
    if '/pdf/' in str(pdf_path):
        docx_dir = Path(str(pdf_path.parent).replace('/pdf/', '/docx/'))
        if docx_dir.exists():
            # Ищем .docx файл с тем же кодом документа
            for docx_file in docx_dir.glob('*.docx'):
                if doc_code and doc_code.upper() in docx_file.stem.upper():
                    return str(docx_file)
            # Или любой .docx файл в папке
            docx_files = list(docx_dir.glob('*.docx'))
            if docx_files:
                return str(docx_files[0])
    
    # Способ 2: В той же папке
    same_dir_docx = pdf_path.with_suffix('.docx')
    if same_dir_docx.exists():
        return str(same_dir_docx)
    
    # Способ 3: Поиск по базовой директории
    if docx_base_dir:
        base = Path(docx_base_dir)
        if base.exists() and doc_code:
            for docx_file in base.rglob('*.docx'):
                if doc_code.upper() in docx_file.stem.upper():
                    return str(docx_file)
    
    return None


if __name__ == "__main__":
    # Тест на примере
    import sys
    
    if len(sys.argv) < 2:
        print("Использование: python docx_parser.py <docx_path> [pdf_path]")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    
    print(f"=== DOCX: {docx_path} ===\n")
    
    # Извлекаем структуру
    headings = extract_structure_docx(docx_path)
    print(f"Найдено заголовков: {len(headings)}\n")
    
    for h in headings:
        print(f"  {'  ' * (h.level - 1)}{h.level}. {h.text[:60]}")
    
    # Если указан PDF - валидируем
    if len(sys.argv) >= 3:
        pdf_path = sys.argv[2]
        print(f"\n=== Валидация vs PDF: {pdf_path} ===\n")
        
        result = validate_docx_vs_pdf(docx_path, pdf_path)
        
        print(f"Результат: {'✅ VALID' if result.is_valid else '❌ INVALID'}")
        print(f"Детали: {result.details}")
