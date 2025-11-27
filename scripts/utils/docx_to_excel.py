#!/usr/bin/env python3
"""
DOCX → XLSX Converter

Назначение:
- Читать DOCX (Word) документ
- Определять заголовок документа и разделы по стилям Heading 1..6
- Для каждого раздела создавать отдельный лист Excel
- Содержимое раздела представлять таблицей:
  - Если в разделе встречаются таблицы Word — выгрузить их построчно
  - Если таблиц нет — выгрузить параграфы как строки (одна колонка "Текст")
- Отдельный лист для заголовка документа (если есть пролог до первого заголовка)

Сохранение:
- output/<base_name>/<base_name>.xlsx

Правила проекта:
- Использовать pathlib.Path
- Кроссплатформенные пути

Требования:
- python-docx
- openpyxl
"""

from __future__ import annotations

import sys
from pathlib import Path
from typing import List, Tuple, Optional

import re
from docx import Document  # type: ignore
from docx.text.paragraph import Paragraph  # type: ignore
from docx.table import Table  # type: ignore
import openpyxl
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.worksheet import Worksheet


def sanitize_sheet_name(name: str) -> str:
    """Привести имя листа к допустимому для Excel (≤31, без запрещённых символов)."""
    invalid = set('[]:*?/\\')
    cleaned = ''.join(ch for ch in name if ch not in invalid and ord(ch) >= 32)
    cleaned = cleaned.strip()
    if not cleaned:
        cleaned = "Sheet"
    if len(cleaned) > 31:
        cleaned = cleaned[:31]
    return cleaned


def docx_iter_blocks(doc: Document) -> List[Tuple[str, object]]:
    """
    Последовательно вернуть список блоков документа, помечая тип:
    - ("p", Paragraph)
    - ("t", Table)
    Порядок соответствует порядку в документе.
    """
    blocks: List[Tuple[str, object]] = []
    # python-docx не даёт прямой итерации смешанных блоков,
    # но можно пройти через document.element.body и сопоставить с API.
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag.endswith('tbl'):
            # Найти соответствующий Table
            for tbl in doc.tables:
                if tbl._element is child:
                    blocks.append(("t", tbl))
                    break
        elif child.tag.endswith('p'):
            for p in doc.paragraphs:
                if p._element is child:
                    blocks.append(("p", p))
                    break
    return blocks


SECTION_HEADING_REGEX = re.compile(
    r"^(?:раздел|подраздел|приложение|section)\s+([0-9ivxlcdm]+)",
    re.IGNORECASE,
)


def is_heading(par: Paragraph) -> Optional[int]:
    """
    Определить уровень заголовка.

    1. Сначала пытаемся распознать стандартные стили Heading 1..6.
    2. Если стиль обычный, но текст начинается с "Раздел I.", "Раздел 1.", "Приложение А" и т.д.,
       считаем это заголовком раздела уровня 1.
    """
    style_name = (par.style.name or "").lower()
    for lvl in range(1, 7):
        if style_name.startswith(f"heading {lvl}") or style_name.startswith(f"заголовок {lvl}"):
            return lvl

    text = par.text.strip()
    if text and SECTION_HEADING_REGEX.match(text):
        return 1

    return None


def split_into_sections(doc: Document) -> List[Tuple[str, List[Tuple[str, object]]]]:
    """
    Разрезать документ на разделы по заголовкам (Heading 1..6).
    Возвращает список кортежей: (заголовок раздела, блоки раздела).
    Если до первого заголовка есть пролог — он станет разделом с именем "Документ".
    """
    blocks = docx_iter_blocks(doc)
    sections: List[Tuple[str, List[Tuple[str, object]]]] = []

    current_title = "Документ"
    current_blocks: List[Tuple[str, object]] = []
    saw_any_heading = False

    for kind, blk in blocks:
        if kind == "p":
            lvl = is_heading(blk)  # type: ignore[arg-type]
            if lvl is not None:
                # Завершить предыдущий раздел
                if current_blocks:
                    sections.append((current_title, current_blocks))
                # Начать новый
                title_text = blk.text.strip() or f"Раздел (Heading {lvl})"
                current_title = title_text
                current_blocks = []
                saw_any_heading = True
                continue
        # Иначе добавляем блок в текущий раздел
        current_blocks.append((kind, blk))

    # Финальный раздел
    if current_blocks:
        sections.append((current_title, current_blocks))

    # Если вообще не было заголовков, сделать один раздел по всему документу
    if not saw_any_heading and not sections:
        sections.append(("Документ", blocks))

    return sections


NUMERIC_ONLY_RE = re.compile(r"^-?\d+([.,]\d+)?$")
DATE_RE = re.compile(r"\d{1,2}[./]\d{1,2}[./]\d{2,4}")
RAW_NUMBER_RE = re.compile(r"^-?[\d.,]+$")
CURRENCY_KEYWORDS = ("сумма", "в валюте", "сальдо")


def _normalize_cell_text(text: str) -> str:
    return " ".join(text.replace("\xa0", " ").split())


def extract_table_matrix(table: Table) -> List[List[str]]:
    """Преобразовать таблицу Word в матрицу строк."""
    matrix: List[List[str]] = []
    for row in table.rows:
        matrix.append([_normalize_cell_text(cell.text) for cell in row.cells])
    return matrix


def is_data_row(row: List[str]) -> bool:
    """
    Эвристично определить, является ли строка строкой данных.
    Считаем строкой данных, если в первом непустом столбце есть число или дата.
    """
    for text in row:
        text = text.strip()
        if not text:
            continue
        if text.lower().startswith("№"):
            return False
        if DATE_RE.search(text):
            return True
        if NUMERIC_ONLY_RE.match(text):
            return True
        if any(ch.isdigit() for ch in text) and "№" not in text:
            return True
        break
    return False


def collapse_header_rows(matrix: List[List[str]], header_rows: int) -> Optional[List[str]]:
    """Объединить многоуровневые заголовки в одну строку."""
    if header_rows <= 0:
        return None
    max_cols = max(len(row) for row in matrix) if matrix else 0
    collapsed: List[str] = []
    for col in range(max_cols):
        parts: List[str] = []
        for idx in range(header_rows):
            row = matrix[idx]
            if col >= len(row):
                continue
            text = row[col].strip()
            if not text:
                continue
            text = text.replace("\n", " ").strip()
            if not any(text.lower() == existing.lower() for existing in parts):
                parts.append(text)
        collapsed.append(" / ".join(parts))
    if all(not part for part in collapsed):
        return None
    return collapsed


def _convert_value(value: str):
    """Преобразовать строку в число, если это похоже на числовое значение."""
    text = value.strip()
    if not text:
        return None

    normalized = (
        text.replace("\xa0", "")
        .replace(" ", "")
        .replace("−", "-")
        .replace("—", "-")
    )

    if RAW_NUMBER_RE.fullmatch(normalized):
        candidate = normalized
        if "," in candidate and "." in candidate:
            if candidate.rfind(".") > candidate.rfind(","):
                candidate = candidate.replace(",", "")
            else:
                candidate = candidate.replace(".", "").replace(",", ".")
        elif "," in candidate:
            if candidate.count(",") == 1:
                candidate = candidate.replace(",", ".")
            else:
                candidate = candidate.replace(",", "")

        try:
            number = float(candidate)
            if number.is_integer():
                return int(number)
            return number
        except ValueError:
            pass

    return text


def write_table(ws: Worksheet, table: Table, start_row: int) -> int:
    """Записать Word-таблицу в Excel, вернуть следующую свободную строку."""
    matrix = extract_table_matrix(table)
    header_rows = 0
    for row in matrix:
        if is_data_row(row):
            break
        header_rows += 1

    r = start_row
    collapsed_header = collapse_header_rows(matrix, header_rows)
    currency_columns: set[int] = set()
    if collapsed_header:
        for col_idx, value in enumerate(collapsed_header, 1):
            ws.cell(row=r, column=col_idx, value=value)
            header_lower = (value or "").lower()
            if any(keyword in header_lower for keyword in CURRENCY_KEYWORDS):
                currency_columns.add(col_idx)
        r += 1

    data_rows = matrix[header_rows:] if header_rows else matrix
    for row in data_rows:
        if not any(cell for cell in row):
            continue
        for col_idx, value in enumerate(row, 1):
            converted = _convert_value(value)
            cell = ws.cell(row=r, column=col_idx, value=converted)
            if col_idx in currency_columns and isinstance(converted, (int, float)):
                cell.number_format = "#,##0.00"
        r += 1
    return r


def write_paragraphs(ws: Worksheet, paragraphs: List[Paragraph], start_row: int) -> int:
    """Записать параграфы построчно в одну колонку 'Текст'."""
    r = start_row
    # Заголовок колонки, если начинаем с первой строки
    if r == 1:
        ws.cell(row=r, column=1, value="Текст")
        r += 1
    for p in paragraphs:
        text = p.text.strip()
        if not text:
            continue
        ws.cell(row=r, column=1, value=text)
        r += 1
    return r


def autosize_columns(ws: Worksheet) -> None:
    """Автоматически подобрать ширину колонок по максимальной длине содержимого."""
    for col_idx in range(1, ws.max_column + 1):
        col_letter = get_column_letter(col_idx)
        max_len = 0
        for row in ws.iter_rows(min_col=col_idx, max_col=col_idx, values_only=True):
            val = row[0]
            if val is None:
                continue
            length = len(str(val))
            if length > max_len:
                max_len = length
        ws.column_dimensions[col_letter].width = min(max(12, max_len + 2), 80)


def convert_docx_to_xlsx(input_path: Path, output_path: Path) -> Path:
    doc = Document(str(input_path))
    sections = split_into_sections(doc)

    wb = openpyxl.Workbook()
    # Удалим дефолтный лист, будем создавать свои
    wb.remove(wb.active)

    # Пролог ("Документ") выведем первым, если присутствует
    for title, blocks in sections:
        sheet_title = sanitize_sheet_name(title)
        # Избежать дубликатов имён листов
        base = sheet_title
        idx = 1
        while sheet_title in wb.sheetnames:
            suffix = f"_{idx}"
            sheet_title = sanitize_sheet_name((base[: (31 - len(suffix))] + suffix))
            idx += 1

        ws = wb.create_sheet(title=sheet_title)

        # Сбор параграфов и таблиц отдельно, чтобы сохранить порядок вывода блоками
        row_cursor = 1
        buffer_pars: List[Paragraph] = []

        def flush_pars() -> None:
            nonlocal row_cursor, buffer_pars
            if buffer_pars:
                row_cursor = write_paragraphs(ws, buffer_pars, row_cursor)
                buffer_pars = []

        for kind, blk in blocks:
            if kind == "p":
                buffer_pars.append(blk)  # type: ignore[arg-type]
            elif kind == "t":
                # Сначала выгружаем накопленные параграфы
                flush_pars()
                row_cursor = write_table(ws, blk, row_cursor)  # type: ignore[arg-type]

        # В конце — оставшиеся параграфы
        flush_pars()

        autosize_columns(ws)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(output_path))
    return output_path


def main() -> None:
    if len(sys.argv) < 2:
        print("Использование: python3 scripts/utils/docx_to_excel.py <путь_к_DOCX> [выходной_XLSX]")
        sys.exit(1)

    in_path = Path(sys.argv[1])
    if not in_path.exists() or in_path.suffix.lower() not in (".docx",):
        print("❌ Укажите существующий DOCX файл")
        sys.exit(1)

    if len(sys.argv) >= 3:
        out_path = Path(sys.argv[2])
    else:
        base_name = in_path.stem
        # Очистка имени по правилам проекта
        # Убираем содержимое в скобках
        if '(' in base_name and ')' in base_name and base_name.index('(') < base_name.index(')'):
            base_name = base_name[: base_name.index('(')].strip()
        # Удаляем спецсимволы, кроме дефиса и подчеркивания
        allowed = set("-_ ")
        cleaned = []
        for ch in base_name:
            if ch.isalnum() or ch in allowed or ch in "АБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯабвгдеёжзийклмнопрстуфхцчшщъыьэюя":
                cleaned.append(ch)
        base_name = ''.join(cleaned)
        base_name = base_name.replace(' ', '_')
        while '__' in base_name:
            base_name = base_name.replace('__', '_')
        base_name = base_name.strip('_')

        out_dir = Path("output") / base_name
        out_path = out_dir / f"{base_name}.xlsx"

    result = convert_docx_to_xlsx(in_path, out_path)
    print(f"✅ Готово: {result}")


if __name__ == "__main__":
    main()
